"""Tests for schema analyzer service."""
import tempfile
import pandas as pd
import pytest
from pathlib import Path

from app.services.schema_analyzer import SchemaAnalyzer


class TestSchemaAnalyzer:
    """Test SchemaAnalyzer service."""

    @pytest.fixture
    def analyzer(self):
        """Create a SchemaAnalyzer instance."""
        return SchemaAnalyzer(sample_rows=1000)

    @pytest.fixture
    def sample_df(self):
        """Create a sample DataFrame for testing."""
        return pd.DataFrame({
            "id": [1, 2, 3, 4, 5],
            "name": ["Alice", "Bob", "Charlie", "Diana", "Eve"],
            "age": [25, 30, 35, 28, 32],
            "salary": [50000.0, 60000.0, 75000.0, 55000.0, 65000.0],
            "department": ["Engineering", "Marketing", "Engineering", "Sales", "Marketing"],
            "is_active": [True, True, False, True, True],
        })

    @pytest.fixture
    def temp_csv_file(self, sample_df):
        """Create a temporary CSV file for testing."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False, newline="") as f:
            sample_df.to_csv(f, index=False)
            temp_path = Path(f.name)
        yield temp_path
        # Cleanup
        if temp_path.exists():
            temp_path.unlink()

    def test_analyze_dataframe_basic(self, analyzer, sample_df):
        """Test basic DataFrame analysis."""
        result = analyzer.analyze_dataframe(sample_df)

        assert result["row_count"] == 5
        assert result["column_count"] == 6
        assert result["sample_rows"] == 5
        assert len(result["columns"]) == 6

    def test_analyze_dataframe_column_names(self, analyzer, sample_df):
        """Test that all column names are captured."""
        result = analyzer.analyze_dataframe(sample_df)

        column_names = [col["name"] for col in result["columns"]]
        assert "id" in column_names
        assert "name" in column_names
        assert "age" in column_names
        assert "salary" in column_names
        assert "department" in column_names
        assert "is_active" in column_names

    def test_analyze_numeric_column(self, analyzer, sample_df):
        """Test numeric column statistics."""
        result = analyzer.analyze_dataframe(sample_df)

        age_col = next(col for col in result["columns"] if col["name"] == "age")
        assert age_col["inferred_type"] == "numeric"
        assert age_col["min"] == 25
        assert age_col["max"] == 35
        assert "mean" in age_col
        assert "median" in age_col
        assert "std" in age_col

    def test_analyze_categorical_column(self, analyzer, sample_df):
        """Test categorical column statistics."""
        result = analyzer.analyze_dataframe(sample_df)

        dept_col = next(col for col in result["columns"] if col["name"] == "department")
        assert dept_col["inferred_type"] == "categorical"
        assert "top_values" in dept_col
        assert "mode" in dept_col
        assert dept_col["unique_count"] == 3

    def test_analyze_boolean_column(self, analyzer, sample_df):
        """Test boolean column detection."""
        result = analyzer.analyze_dataframe(sample_df)

        active_col = next(col for col in result["columns"] if col["name"] == "is_active")
        assert active_col["inferred_type"] == "boolean"
        assert active_col["unique_count"] == 2

    def test_analyze_csv_file(self, analyzer, temp_csv_file):
        """Test analyzing a CSV file directly."""
        result = analyzer.analyze_csv(temp_csv_file)

        assert result["row_count"] == 5
        assert result["column_count"] == 6
        assert len(result["columns"]) == 6

    def test_analyze_csv_with_delimiter(self, analyzer):
        """Test analyzing CSV with custom delimiter."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
            f.write("name;value;count\n")
            f.write("item1;100;5\n")
            f.write("item2;200;10\n")
            csv_path = Path(f.name)

        result = analyzer.analyze_csv(csv_path, delimiter=";")

        assert result["column_count"] == 3
        column_names = [col["name"] for col in result["columns"]]
        assert "name" in column_names
        assert "value" in column_names
        assert "count" in column_names

    def test_analyze_null_values(self, analyzer):
        """Test null value detection."""
        df = pd.DataFrame({
            "complete": [1, 2, 3, 4, 5],
            "with_nulls": [1, None, 3, None, 5],
        })

        result = analyzer.analyze_dataframe(df)

        complete_col = next(col for col in result["columns"] if col["name"] == "complete")
        assert complete_col["null_count"] == 0
        assert complete_col["null_percentage"] == 0.0

        nulls_col = next(col for col in result["columns"] if col["name"] == "with_nulls")
        assert nulls_col["null_count"] == 2
        assert nulls_col["null_percentage"] == 40.0

    def test_analyze_datetime_column(self, analyzer):
        """Test datetime column detection."""
        df = pd.DataFrame({
            "date_str": ["2024-01-01", "2024-02-15", "2024-03-20"],
            "value": [1, 2, 3],
        })

        result = analyzer.analyze_dataframe(df)

        date_col = next(col for col in result["columns"] if col["name"] == "date_str")
        assert date_col["inferred_type"] == "datetime"
        assert "min" in date_col
        assert "max" in date_col

    def test_analyze_text_column(self, analyzer):
        """Test text column detection (high cardinality strings)."""
        # Create DataFrame with unique text values
        df = pd.DataFrame({
            "description": [f"Unique description number {i}" for i in range(100)],
            "id": list(range(100)),
        })

        result = analyzer.analyze_dataframe(df)

        desc_col = next(col for col in result["columns"] if col["name"] == "description")
        assert desc_col["inferred_type"] == "text"

    def test_analyze_file_not_found(self, analyzer):
        """Test that FileNotFoundError is raised for missing files."""
        with pytest.raises(FileNotFoundError):
            analyzer.analyze_csv("/nonexistent/path/file.csv")

    def test_empty_dataframe(self, analyzer):
        """Test analyzing an empty DataFrame."""
        df = pd.DataFrame({"col1": [], "col2": []})

        result = analyzer.analyze_dataframe(df)

        assert result["row_count"] == 0
        assert result["column_count"] == 2

    def test_unique_count(self, analyzer):
        """Test unique value counting."""
        df = pd.DataFrame({
            "all_same": [1, 1, 1, 1, 1],
            "all_unique": [1, 2, 3, 4, 5],
        })

        result = analyzer.analyze_dataframe(df)

        same_col = next(col for col in result["columns"] if col["name"] == "all_same")
        assert same_col["unique_count"] == 1

        unique_col = next(col for col in result["columns"] if col["name"] == "all_unique")
        assert unique_col["unique_count"] == 5


class TestSchemaAnalyzerFileTypes:
    """Test SchemaAnalyzer with different file types."""

    @pytest.fixture
    def analyzer(self):
        """Create a SchemaAnalyzer instance."""
        return SchemaAnalyzer(sample_rows=1000)

    @pytest.fixture
    def sample_df(self):
        """Create a sample DataFrame for testing."""
        return pd.DataFrame({
            "id": [1, 2, 3],
            "name": ["Alice", "Bob", "Charlie"],
            "value": [100.5, 200.75, 300.25],
        })

    def test_analyze_excel_file(self, analyzer, sample_df):
        """Test analyzing an Excel file."""
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            excel_path = Path(f.name)
        sample_df.to_excel(excel_path, index=False)

        try:
            result = analyzer.analyze_excel(excel_path)

            assert result["file_type"] == "excel"
            assert result["row_count"] == 3
            assert result["column_count"] == 3
            assert "sheet_names" in result
            assert "analyzed_sheet" in result

            column_names = [col["name"] for col in result["columns"]]
            assert "id" in column_names
            assert "name" in column_names
            assert "value" in column_names
        finally:
            excel_path.unlink()

    def test_analyze_excel_specific_sheet(self, analyzer):
        """Test analyzing a specific Excel sheet."""
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            excel_path = Path(f.name)

        # Create Excel with multiple sheets
        with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
            pd.DataFrame({"sheet1_col": [1, 2, 3]}).to_excel(writer, sheet_name="Sheet1", index=False)
            pd.DataFrame({"sheet2_col": [4, 5, 6]}).to_excel(writer, sheet_name="Sheet2", index=False)

        try:
            result = analyzer.analyze_excel(excel_path, sheet_name="Sheet2")

            assert result["analyzed_sheet"] == "Sheet2"
            column_names = [col["name"] for col in result["columns"]]
            assert "sheet2_col" in column_names
        finally:
            excel_path.unlink()

    def test_analyze_json_file(self, analyzer, sample_df):
        """Test analyzing a JSON file."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as f:
            sample_df.to_json(f.name, orient="records")
            json_path = Path(f.name)

        try:
            result = analyzer.analyze_json(json_path)

            assert result["file_type"] == "json"
            assert result["row_count"] == 3
            assert result["column_count"] == 3

            column_names = [col["name"] for col in result["columns"]]
            assert "id" in column_names
            assert "name" in column_names
            assert "value" in column_names
        finally:
            json_path.unlink()

    def test_analyze_parquet_file(self, analyzer, sample_df):
        """Test analyzing a Parquet file."""
        with tempfile.NamedTemporaryFile(suffix=".parquet", delete=False) as f:
            parquet_path = Path(f.name)
        sample_df.to_parquet(parquet_path, index=False)

        try:
            result = analyzer.analyze_parquet(parquet_path)

            assert result["file_type"] == "parquet"
            assert result["row_count"] == 3
            assert result["column_count"] == 3

            column_names = [col["name"] for col in result["columns"]]
            assert "id" in column_names
            assert "name" in column_names
            assert "value" in column_names
        finally:
            parquet_path.unlink()

    def test_analyze_text_file_plain(self, analyzer):
        """Test analyzing a plain text file."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("Line one\n")
            f.write("Line two\n")
            f.write("Line three\n")
            text_path = Path(f.name)

        try:
            result = analyzer.analyze_text(text_path)

            assert result["file_type"] == "text"
            assert result["row_count"] == 3
            assert result["column_count"] == 2
            assert "text_stats" in result
            assert result["text_stats"]["total_lines"] == 3

            column_names = [col["name"] for col in result["columns"]]
            assert "line_number" in column_names
            assert "content" in column_names
        finally:
            text_path.unlink()

    def test_analyze_text_file_delimited(self, analyzer):
        """Test analyzing a text file that contains delimited data."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("name,age,city\n")
            f.write("Alice,25,NYC\n")
            f.write("Bob,30,LA\n")
            text_path = Path(f.name)

        try:
            result = analyzer.analyze_text(text_path)

            assert result["file_type"] == "text"
            # Should detect as delimited and parse
            assert result["column_count"] == 3
            assert "detected_format" in result
            assert "delimited" in result["detected_format"]

            column_names = [col["name"] for col in result["columns"]]
            assert "name" in column_names
            assert "age" in column_names
            assert "city" in column_names
        finally:
            text_path.unlink()

    def test_analyze_word_file_with_table(self, analyzer):
        """Test analyzing a Word document with a table."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            word_path = Path(f.name)

        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        # Header row
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Age"
        table.rows[0].cells[2].text = "City"
        # Data row 1
        table.rows[1].cells[0].text = "Alice"
        table.rows[1].cells[1].text = "25"
        table.rows[1].cells[2].text = "NYC"
        # Data row 2
        table.rows[2].cells[0].text = "Bob"
        table.rows[2].cells[1].text = "30"
        table.rows[2].cells[2].text = "LA"
        doc.save(word_path)

        try:
            result = analyzer.analyze_word(word_path)

            assert result["file_type"] == "word"
            assert result["has_tables"] is True
            assert result["table_count"] == 1
            assert result["row_count"] == 2  # Data rows, not header

            column_names = [col["name"] for col in result["columns"]]
            assert "Name" in column_names
            assert "Age" in column_names
            assert "City" in column_names
        finally:
            word_path.unlink()

    def test_analyze_word_file_paragraphs_only(self, analyzer):
        """Test analyzing a Word document with only paragraphs."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            word_path = Path(f.name)

        doc = Document()
        doc.add_paragraph("First paragraph of text.")
        doc.add_paragraph("Second paragraph of text.")
        doc.add_paragraph("Third paragraph of text.")
        doc.save(word_path)

        try:
            result = analyzer.analyze_word(word_path)

            assert result["file_type"] == "word"
            assert result["row_count"] == 3  # 3 paragraphs
            assert result["column_count"] == 2
            assert "text_stats" in result
            assert result["text_stats"]["total_paragraphs"] == 3
            assert result["text_stats"]["has_tables"] is False

            column_names = [col["name"] for col in result["columns"]]
            assert "paragraph_number" in column_names
            assert "content" in column_names
        finally:
            word_path.unlink()

    def test_analyze_file_auto_detect_csv(self, analyzer):
        """Test analyze_file auto-detects CSV files."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
            f.write("col1,col2\n1,2\n3,4\n")
            csv_path = Path(f.name)

        try:
            result = analyzer.analyze_file(csv_path)
            assert result["file_type"] == "csv"
        finally:
            csv_path.unlink()

    def test_analyze_file_auto_detect_excel(self, analyzer, sample_df):
        """Test analyze_file auto-detects Excel files."""
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            excel_path = Path(f.name)
        sample_df.to_excel(excel_path, index=False)

        try:
            result = analyzer.analyze_file(excel_path)
            assert result["file_type"] == "excel"
        finally:
            excel_path.unlink()

    def test_analyze_file_unsupported_type(self, analyzer):
        """Test that unsupported file types raise ValueError."""
        with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
            unsupported_path = Path(f.name)

        try:
            with pytest.raises(ValueError, match="Unsupported file type"):
                analyzer.analyze_file(unsupported_path)
        finally:
            unsupported_path.unlink()
