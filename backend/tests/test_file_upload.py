"""Tests for file upload functionality."""
import io
import os
import tempfile
import pytest
from pathlib import Path
from unittest.mock import patch


class TestFileUpload:
    """Test file upload endpoint."""

    @pytest.fixture
    def project_id(self, client):
        """Create a project and return its ID."""
        response = client.post("/projects", json={"name": "Upload Test Project"})
        return response.json()["id"]

    @pytest.fixture
    def sample_csv_content(self):
        """Sample CSV content for testing."""
        return b"name,age,salary,department\nAlice,30,75000,Engineering\nBob,25,60000,Marketing\nCharlie,35,90000,Engineering\n"

    @pytest.fixture
    def temp_upload_dir(self):
        """Create a temporary upload directory."""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield tmpdir

    def test_upload_csv_file(self, client, project_id, sample_csv_content, temp_upload_dir):
        """Test uploading a CSV file creates data source with schema."""
        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("test_data.csv", io.BytesIO(sample_csv_content), "text/csv")},
                data={"delimiter": ","},
            )

        assert response.status_code == 201
        data = response.json()
        assert data["name"] == "test_data.csv"
        assert data["type"] == "file_upload"
        assert data["project_id"] == project_id
        assert "config_json" in data
        assert data["config_json"]["delimiter"] == ","
        assert data["config_json"]["original_filename"] == "test_data.csv"

        # Check schema was analyzed
        assert "schema_summary" in data
        schema = data["schema_summary"]
        assert schema["row_count"] == 3
        assert schema["column_count"] == 4

    def test_upload_csv_with_custom_name(self, client, project_id, sample_csv_content, temp_upload_dir):
        """Test uploading with a custom name."""
        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("original.csv", io.BytesIO(sample_csv_content), "text/csv")},
                data={"name": "My Custom Dataset", "delimiter": ","},
            )

        assert response.status_code == 201
        data = response.json()
        assert data["name"] == "My Custom Dataset"
        assert data["config_json"]["original_filename"] == "original.csv"

    def test_upload_unsupported_file_type_fails(self, client, project_id, temp_upload_dir):
        """Test that unsupported file types are rejected."""
        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("data.exe", io.BytesIO(b"fake binary"), "application/octet-stream")},
            )

        assert response.status_code == 400
        assert "Unsupported file type" in response.json()["detail"]

    def test_upload_to_nonexistent_project_fails(self, client, sample_csv_content, temp_upload_dir):
        """Test that uploading to non-existent project fails."""
        fake_project_id = "00000000-0000-0000-0000-000000000000"

        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{fake_project_id}/data-sources/upload",
                files={"file": ("test.csv", io.BytesIO(sample_csv_content), "text/csv")},
            )

        assert response.status_code == 404

    def test_upload_schema_analysis_columns(self, client, project_id, temp_upload_dir):
        """Test that schema analysis correctly identifies column types."""
        csv_content = b"id,name,price,is_active,created_at\n1,Widget,19.99,true,2024-01-15\n2,Gadget,29.99,false,2024-02-20\n"

        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("products.csv", io.BytesIO(csv_content), "text/csv")},
            )

        assert response.status_code == 201
        schema = response.json()["schema_summary"]

        # Check columns were analyzed
        columns_by_name = {col["name"]: col for col in schema["columns"]}

        assert "id" in columns_by_name
        assert "name" in columns_by_name
        assert "price" in columns_by_name

        # Check numeric column has stats
        price_col = columns_by_name["price"]
        assert price_col["inferred_type"] == "numeric"
        assert "min" in price_col
        assert "max" in price_col
        assert price_col["min"] == 19.99
        assert price_col["max"] == 29.99

    def test_upload_with_semicolon_delimiter(self, client, project_id, temp_upload_dir):
        """Test uploading CSV with semicolon delimiter."""
        csv_content = b"name;value\ntest;123\n"

        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("data.csv", io.BytesIO(csv_content), "text/csv")},
                data={"delimiter": ";"},
            )

        assert response.status_code == 201
        schema = response.json()["schema_summary"]
        assert schema["column_count"] == 2
        column_names = [col["name"] for col in schema["columns"]]
        assert "name" in column_names
        assert "value" in column_names


class TestFileUploadMultipleTypes:
    """Test file upload with various file types."""

    @pytest.fixture
    def project_id(self, client):
        """Create a project and return its ID."""
        response = client.post("/projects", json={"name": "Multi-type Upload Project"})
        return response.json()["id"]

    @pytest.fixture
    def temp_upload_dir(self):
        """Create a temporary upload directory."""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield tmpdir

    def test_upload_json_file(self, client, project_id, temp_upload_dir):
        """Test uploading a JSON file."""
        json_content = b'[{"id": 1, "name": "Alice"}, {"id": 2, "name": "Bob"}]'

        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("data.json", io.BytesIO(json_content), "application/json")},
            )

        assert response.status_code == 201
        data = response.json()
        assert data["config_json"]["file_type"] == "json"
        schema = data["schema_summary"]
        assert schema["file_type"] == "json"
        assert schema["row_count"] == 2
        column_names = [col["name"] for col in schema["columns"]]
        assert "id" in column_names
        assert "name" in column_names

    def test_upload_text_file(self, client, project_id, temp_upload_dir):
        """Test uploading a text file."""
        text_content = b"Line one\nLine two\nLine three\n"

        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("notes.txt", io.BytesIO(text_content), "text/plain")},
            )

        assert response.status_code == 201
        data = response.json()
        assert data["config_json"]["file_type"] == "text"
        schema = data["schema_summary"]
        assert schema["file_type"] == "text"
        assert schema["row_count"] == 3
        assert "text_stats" in schema

    def test_upload_excel_file(self, client, project_id, temp_upload_dir):
        """Test uploading an Excel file."""
        import pandas as pd

        # Create Excel content in memory
        df = pd.DataFrame({"name": ["Alice", "Bob"], "value": [100, 200]})
        excel_buffer = io.BytesIO()
        df.to_excel(excel_buffer, index=False, engine="openpyxl")
        excel_buffer.seek(0)

        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("data.xlsx", excel_buffer, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            )

        assert response.status_code == 201
        data = response.json()
        assert data["config_json"]["file_type"] == "excel"
        schema = data["schema_summary"]
        assert schema["file_type"] == "excel"
        assert schema["row_count"] == 2
        assert "sheet_names" in schema
        column_names = [col["name"] for col in schema["columns"]]
        assert "name" in column_names
        assert "value" in column_names

    def test_upload_parquet_file(self, client, project_id, temp_upload_dir):
        """Test uploading a Parquet file."""
        import pandas as pd

        # Create Parquet content in memory
        df = pd.DataFrame({"id": [1, 2, 3], "score": [85.5, 90.0, 78.5]})
        parquet_buffer = io.BytesIO()
        df.to_parquet(parquet_buffer, index=False)
        parquet_buffer.seek(0)

        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("data.parquet", parquet_buffer, "application/octet-stream")},
            )

        assert response.status_code == 201
        data = response.json()
        assert data["config_json"]["file_type"] == "parquet"
        schema = data["schema_summary"]
        assert schema["file_type"] == "parquet"
        assert schema["row_count"] == 3
        column_names = [col["name"] for col in schema["columns"]]
        assert "id" in column_names
        assert "score" in column_names

    def test_upload_word_file(self, client, project_id, temp_upload_dir):
        """Test uploading a Word document."""
        from docx import Document

        # Create Word document in memory
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        word_buffer = io.BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)

        with patch("app.api.data_sources.get_settings") as mock_settings:
            mock_settings.return_value.upload_dir = temp_upload_dir
            mock_settings.return_value.max_upload_size_mb = 100

            response = client.post(
                f"/projects/{project_id}/data-sources/upload",
                files={"file": ("document.docx", word_buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )

        assert response.status_code == 201
        data = response.json()
        assert data["config_json"]["file_type"] == "word"
        schema = data["schema_summary"]
        assert schema["file_type"] == "word"
        assert schema["row_count"] == 2  # 2 paragraphs
