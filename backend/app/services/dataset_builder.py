"""Dataset builder service for constructing DataFrames from DatasetSpecs."""
import pandas as pd
from pathlib import Path
from typing import Any
from uuid import UUID

from sqlalchemy.orm import Session

from app.models.dataset_spec import DatasetSpec
from app.models.data_source import DataSource


class DatasetBuilder:
    """Builds DataFrames from DatasetSpec definitions."""

    def __init__(self, db: Session):
        """Initialize builder with database session.

        Args:
            db: SQLAlchemy database session
        """
        self.db = db

    def build_dataset_from_spec(
        self,
        dataset_spec_id: UUID,
        sample_size: int | None = None,
    ) -> pd.DataFrame:
        """Build a DataFrame from a DatasetSpec.

        Args:
            dataset_spec_id: UUID of the DatasetSpec to build
            sample_size: Optional limit on rows to load (for validation/preview)

        Returns:
            Pandas DataFrame with selected features and target

        Raises:
            ValueError: If DatasetSpec not found or has invalid configuration
        """
        # Get the dataset spec
        spec = self.db.query(DatasetSpec).filter(DatasetSpec.id == dataset_spec_id).first()
        if not spec:
            raise ValueError(f"DatasetSpec {dataset_spec_id} not found")

        # Get data sources from spec
        # Handle both list format (direct list of IDs) and dict format ({"sources": [...]})
        data_sources_config = spec.data_sources_json or {}
        if isinstance(data_sources_config, list):
            # Direct list of source IDs
            source_ids = data_sources_config
        else:
            # Dict format - try various keys
            source_ids = data_sources_config.get("sources", [])
            if not source_ids:
                source_ids = data_sources_config.get("source_ids", [])
            if not source_ids:
                # Try single source_id for backward compatibility
                source_id = data_sources_config.get("source_id")
                if source_id:
                    source_ids = [source_id]
            if not source_ids:
                # Try "primary" key (used by frontend)
                primary_id = data_sources_config.get("primary")
                if primary_id:
                    source_ids = [primary_id]

        if not source_ids:
            raise ValueError(f"DatasetSpec {dataset_spec_id} has no data sources configured")

        # Load and concatenate data from all sources
        dfs = []
        for source_entry in source_ids:
            # Handle both dict format {'source_id': '...', 'table_name': '...'} and direct ID
            if isinstance(source_entry, dict):
                source_id = source_entry.get("source_id") or source_entry.get("id")
            else:
                source_id = source_entry

            if source_id:
                df = self._load_data_source(source_id)
                dfs.append(df)

        # Concatenate if multiple sources
        if len(dfs) == 1:
            combined_df = dfs[0]
        else:
            combined_df = pd.concat(dfs, ignore_index=True)

        # Apply sample_size limit if specified (for validation/preview)
        if sample_size is not None and len(combined_df) > sample_size:
            combined_df = combined_df.head(sample_size)

        # Apply feature engineering if specified in spec_json
        spec_json = spec.spec_json or {}
        engineered_features = spec_json.get("engineered_features", [])
        if engineered_features:
            from app.services.feature_engineering import apply_feature_engineering
            # Use strict=False to skip failed features instead of crashing
            combined_df = apply_feature_engineering(combined_df, engineered_features, inplace=False, strict=False)

        # Create target column if specified and doesn't exist
        target_creation = spec_json.get("target_creation")
        target_column = spec.target_column or spec_json.get("target_column")

        if target_creation and target_column and target_column not in combined_df.columns:
            from app.services.feature_engineering import apply_target_creation
            # Ensure target_creation has the column_name set
            if "column_name" not in target_creation:
                target_creation = {**target_creation, "column_name": target_column}
            combined_df = apply_target_creation(combined_df, target_creation, inplace=False)

            # Drop rows with NaN target (common for shifted targets)
            combined_df = combined_df.dropna(subset=[target_column])

        # Apply column selection
        combined_df = self._select_columns(combined_df, spec)

        # Apply filters if any
        combined_df = self._apply_filters(combined_df, spec)

        return combined_df

    def _load_data_source(self, source_id: str | UUID) -> pd.DataFrame:
        """Load data from a data source.

        Args:
            source_id: UUID of the data source

        Returns:
            Pandas DataFrame with the source data

        Raises:
            ValueError: If data source not found or unsupported type
        """
        # Convert string to UUID if needed
        if isinstance(source_id, str):
            source_id = UUID(source_id)

        source = self.db.query(DataSource).filter(DataSource.id == source_id).first()
        if not source:
            raise ValueError(f"DataSource {source_id} not found")

        if source.type == "file_upload":
            return self._load_file_upload(source)
        else:
            raise ValueError(f"Unsupported data source type: {source.type}")

    def _load_file_upload(self, source: DataSource) -> pd.DataFrame:
        """Load data from a file upload data source.

        Supports: CSV, Excel, JSON, Parquet, Text, Word

        Args:
            source: DataSource model instance

        Returns:
            Pandas DataFrame with the file data

        Raises:
            ValueError: If file not found or invalid configuration
        """
        from app.services.file_storage import ensure_file_on_disk

        config = source.config_json or {}

        path = ensure_file_on_disk(source)

        # Determine file type from config or extension
        file_type = config.get("file_type")
        if not file_type:
            # Fallback to extension detection
            ext = path.suffix.lower()
            file_type_map = {
                ".csv": "csv",
                ".xlsx": "excel",
                ".xls": "excel",
                ".json": "json",
                ".parquet": "parquet",
                ".txt": "text",
                ".docx": "word",
                ".doc": "word",
            }
            file_type = file_type_map.get(ext, "csv")

        # Load based on file type
        if file_type == "csv":
            delimiter = config.get("delimiter", ",")
            return pd.read_csv(path, delimiter=delimiter)

        elif file_type == "excel":
            sheet_name = config.get("sheet_name", 0)
            return pd.read_excel(path, sheet_name=sheet_name)

        elif file_type == "json":
            return pd.read_json(path)

        elif file_type == "parquet":
            return pd.read_parquet(path)

        elif file_type == "text":
            # Try to detect delimiter first
            with open(path, "r", encoding="utf-8") as f:
                first_line = f.readline()

            for delimiter in [",", "\t", ";", "|"]:
                if delimiter in first_line:
                    try:
                        df = pd.read_csv(path, delimiter=delimiter)
                        if len(df.columns) > 1:
                            return df
                    except Exception:
                        pass

            # Plain text - return as lines
            with open(path, "r", encoding="utf-8") as f:
                lines = f.read().strip().split("\n")
            return pd.DataFrame({
                "line_number": range(1, len(lines) + 1),
                "content": lines
            })

        elif file_type == "word":
            from docx import Document
            doc = Document(path)

            # Try to extract tables first
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                if len(rows) > 1:
                    headers = rows[0]
                    data = rows[1:]
                    return pd.DataFrame(data, columns=headers)

            # No tables - return paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return pd.DataFrame({
                "paragraph_number": range(1, len(paragraphs) + 1),
                "content": paragraphs
            })

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _select_columns(self, df: pd.DataFrame, spec: DatasetSpec) -> pd.DataFrame:
        """Select columns based on DatasetSpec configuration.

        Args:
            df: Input DataFrame
            spec: DatasetSpec with column configuration

        Returns:
            DataFrame with selected columns
        """
        columns_to_select = []

        # Add feature columns
        feature_columns = spec.feature_columns or []
        for col in feature_columns:
            if col in df.columns:
                columns_to_select.append(col)

        # Add target column
        if spec.target_column and spec.target_column in df.columns:
            if spec.target_column not in columns_to_select:
                columns_to_select.append(spec.target_column)

        # If no columns specified, return all columns
        if not columns_to_select:
            return df

        return df[columns_to_select]

    def _apply_filters(self, df: pd.DataFrame, spec: DatasetSpec) -> pd.DataFrame:
        """Apply filters to DataFrame based on DatasetSpec configuration.

        Args:
            df: Input DataFrame
            spec: DatasetSpec with filter configuration

        Returns:
            Filtered DataFrame
        """
        filters = spec.filters_json or {}

        for column, filter_config in filters.items():
            if column not in df.columns:
                continue

            if isinstance(filter_config, dict):
                # Range filters
                if "min" in filter_config:
                    df = df[df[column] >= filter_config["min"]]
                if "max" in filter_config:
                    df = df[df[column] <= filter_config["max"]]
                # Value list filter
                if "in" in filter_config:
                    df = df[df[column].isin(filter_config["in"])]
                if "not_in" in filter_config:
                    df = df[~df[column].isin(filter_config["not_in"])]
            elif isinstance(filter_config, list):
                # Simple value list filter
                df = df[df[column].isin(filter_config)]
            else:
                # Exact value match
                df = df[df[column] == filter_config]

        return df

    def get_dataset_info(self, dataset_spec_id: UUID) -> dict[str, Any]:
        """Get information about a built dataset without loading all data.

        Args:
            dataset_spec_id: UUID of the DatasetSpec

        Returns:
            Dictionary with dataset information
        """
        spec = self.db.query(DatasetSpec).filter(DatasetSpec.id == dataset_spec_id).first()
        if not spec:
            raise ValueError(f"DatasetSpec {dataset_spec_id} not found")

        # Handle both list format and dict format
        data_sources_config = spec.data_sources_json or {}
        if isinstance(data_sources_config, list):
            source_ids = data_sources_config
        else:
            source_ids = data_sources_config.get("sources", [])
            if not source_ids:
                source_ids = data_sources_config.get("source_ids", [])
            if not source_ids:
                source_id = data_sources_config.get("source_id")
                if source_id:
                    source_ids = [source_id]
            if not source_ids:
                primary_id = data_sources_config.get("primary")
                if primary_id:
                    source_ids = [primary_id]

        sources_info = []
        for source_entry in source_ids:
            # Handle both dict format {'source_id': '...', 'table_name': '...'} and direct ID
            if isinstance(source_entry, dict):
                source_id = source_entry.get("source_id") or source_entry.get("id")
            else:
                source_id = source_entry

            if not source_id:
                continue

            if isinstance(source_id, str):
                source_id = UUID(source_id)
            source = self.db.query(DataSource).filter(DataSource.id == source_id).first()
            if source:
                sources_info.append({
                    "id": str(source.id),
                    "name": source.name,
                    "type": source.type,
                    "schema_summary": source.schema_summary,
                })

        return {
            "dataset_spec_id": str(spec.id),
            "name": spec.name,
            "target_column": spec.target_column,
            "feature_columns": spec.feature_columns,
            "filters": spec.filters_json,
            "data_sources": sources_info,
        }
