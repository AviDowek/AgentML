"""Extensible file type handlers for data ingestion.

This module provides a handler-based architecture for reading various file formats
and converting them to pandas DataFrames. New file types can be added by implementing
the FileHandler protocol and registering them.
"""
import logging
import re
from abc import ABC, abstractmethod
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any, Optional, Protocol, Union

import pandas as pd

logger = logging.getLogger(__name__)


class FileHandler(Protocol):
    """Protocol for file type handlers."""

    extensions: list[str]
    file_type: str
    description: str

    def can_handle(self, file_path: Path, content: Optional[bytes] = None) -> bool:
        """Check if this handler can process the given file."""
        ...

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        """Read the file and return (DataFrame, metadata)."""
        ...

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        """Get total row count without loading full file if possible."""
        ...


class BaseFileHandler(ABC):
    """Base class for file handlers with common functionality."""

    extensions: list[str] = []
    file_type: str = "unknown"
    description: str = "Unknown file type"

    def can_handle(self, file_path: Path, content: Optional[bytes] = None) -> bool:
        """Check by extension. Override for content-based detection."""
        return file_path.suffix.lower() in self.extensions

    @abstractmethod
    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        """Read file and return DataFrame with metadata."""
        pass

    @abstractmethod
    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        """Get total row count."""
        pass


# ==================== CSV Handler ====================

class CSVHandler(BaseFileHandler):
    """Handler for CSV and similar delimited files."""

    extensions = [".csv", ".tsv"]
    file_type = "csv"
    description = "Comma/Tab-Separated Values"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        delimiter = kwargs.get("delimiter", ",")
        encoding = kwargs.get("encoding", "utf-8")

        try:
            df = pd.read_csv(
                file_path,
                delimiter=delimiter,
                nrows=sample_rows,
                encoding=encoding,
                on_bad_lines="skip"
            )
        except UnicodeDecodeError:
            # Try with latin-1 fallback
            df = pd.read_csv(
                file_path,
                delimiter=delimiter,
                nrows=sample_rows,
                encoding="latin-1",
                on_bad_lines="skip"
            )
            encoding = "latin-1"

        metadata = {
            "delimiter": delimiter,
            "encoding": encoding,
        }
        return df, metadata

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        encoding = kwargs.get("encoding", "utf-8")
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return sum(1 for _ in f) - 1  # -1 for header
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return sum(1 for _ in f) - 1


# ==================== Excel Handler ====================

class ExcelHandler(BaseFileHandler):
    """Handler for Excel files (.xlsx, .xls)."""

    extensions = [".xlsx", ".xls", ".xlsm", ".xlsb"]
    file_type = "excel"
    description = "Microsoft Excel Spreadsheet"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        sheet_name = kwargs.get("sheet_name", 0)

        with pd.ExcelFile(file_path) as xlsx:
            sheet_names = xlsx.sheet_names
            df = pd.read_excel(xlsx, sheet_name=sheet_name, nrows=sample_rows)

            actual_sheet = sheet_name if isinstance(sheet_name, str) else sheet_names[sheet_name]

        metadata = {
            "sheet_names": sheet_names,
            "analyzed_sheet": actual_sheet,
        }
        return df, metadata

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        sheet_name = kwargs.get("sheet_name", 0)
        with pd.ExcelFile(file_path) as xlsx:
            df = pd.read_excel(xlsx, sheet_name=sheet_name)
            return len(df)


# ==================== JSON Handler ====================

class JSONHandler(BaseFileHandler):
    """Handler for JSON files (arrays, objects, JSON Lines)."""

    extensions = [".json", ".jsonl", ".ndjson"]
    file_type = "json"
    description = "JSON / JSON Lines"

    def can_handle(self, file_path: Path, content: Optional[bytes] = None) -> bool:
        if file_path.suffix.lower() in self.extensions:
            return True
        # Content-based detection
        if content:
            try:
                text = content[:1000].decode("utf-8").strip()
                return text.startswith(("{", "["))
            except Exception:
                pass
        return False

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        # Check if it's JSON Lines format
        with open(file_path, "r", encoding="utf-8") as f:
            first_line = f.readline().strip()

        if first_line.startswith("{") and not first_line.endswith("]"):
            # JSON Lines format
            df = pd.read_json(file_path, lines=True, nrows=sample_rows)
            metadata = {"json_format": "jsonl"}
        else:
            # Standard JSON
            df = pd.read_json(file_path)
            if sample_rows and len(df) > sample_rows:
                df = df.head(sample_rows)
            metadata = {"json_format": "standard"}

        return df, metadata

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        with open(file_path, "r", encoding="utf-8") as f:
            first_line = f.readline().strip()

        if first_line.startswith("{") and not first_line.endswith("]"):
            # JSON Lines - count lines
            with open(file_path, "r", encoding="utf-8") as f:
                return sum(1 for _ in f)
        else:
            df = pd.read_json(file_path)
            return len(df)


# ==================== Parquet Handler ====================

class ParquetHandler(BaseFileHandler):
    """Handler for Apache Parquet files."""

    extensions = [".parquet", ".pq"]
    file_type = "parquet"
    description = "Apache Parquet"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        df = pd.read_parquet(file_path)

        if sample_rows and len(df) > sample_rows:
            df = df.head(sample_rows)

        # Get parquet metadata
        import pyarrow.parquet as pq
        pf = pq.ParquetFile(file_path)
        metadata = {
            "num_row_groups": pf.metadata.num_row_groups,
            "compression": str(pf.schema_arrow),
        }
        return df, metadata

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        import pyarrow.parquet as pq
        pf = pq.ParquetFile(file_path)
        return pf.metadata.num_rows


# ==================== SQLite Handler ====================

class SQLiteHandler(BaseFileHandler):
    """Handler for SQLite database files."""

    extensions = [".sqlite", ".sqlite3", ".db"]
    file_type = "sqlite"
    description = "SQLite Database"

    def can_handle(self, file_path: Path, content: Optional[bytes] = None) -> bool:
        if file_path.suffix.lower() in self.extensions:
            return True
        # Content-based detection - SQLite magic bytes
        if content and content[:16].startswith(b"SQLite format 3"):
            return True
        return False

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        import sqlite3

        table_name = kwargs.get("table_name")

        conn = sqlite3.connect(file_path)
        try:
            # Get list of tables
            tables = pd.read_sql_query(
                "SELECT name FROM sqlite_master WHERE type='table'",
                conn
            )["name"].tolist()

            if not tables:
                raise ValueError("No tables found in SQLite database")

            # Use specified table or first one
            if table_name and table_name in tables:
                target_table = table_name
            else:
                target_table = tables[0]

            # Read the table
            limit_clause = f" LIMIT {sample_rows}" if sample_rows else ""
            df = pd.read_sql_query(
                f'SELECT * FROM "{target_table}"{limit_clause}',
                conn
            )

            metadata = {
                "tables": tables,
                "analyzed_table": target_table,
                "table_count": len(tables),
            }
        finally:
            conn.close()

        return df, metadata

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        import sqlite3

        table_name = kwargs.get("table_name")

        conn = sqlite3.connect(file_path)
        try:
            tables = pd.read_sql_query(
                "SELECT name FROM sqlite_master WHERE type='table'",
                conn
            )["name"].tolist()

            target_table = table_name if table_name in tables else tables[0]
            result = conn.execute(f'SELECT COUNT(*) FROM "{target_table}"').fetchone()
            return result[0]
        finally:
            conn.close()


# ==================== Feather Handler ====================

class FeatherHandler(BaseFileHandler):
    """Handler for Apache Feather/Arrow IPC files."""

    extensions = [".feather", ".arrow", ".ipc"]
    file_type = "feather"
    description = "Apache Feather/Arrow IPC"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        df = pd.read_feather(file_path)

        if sample_rows and len(df) > sample_rows:
            df = df.head(sample_rows)

        return df, {}

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        df = pd.read_feather(file_path)
        return len(df)


# ==================== HDF5 Handler ====================

class HDF5Handler(BaseFileHandler):
    """Handler for HDF5 files."""

    extensions = [".h5", ".hdf5", ".hdf"]
    file_type = "hdf5"
    description = "HDF5 (Hierarchical Data Format)"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        key = kwargs.get("key")

        with pd.HDFStore(file_path, mode="r") as store:
            keys = store.keys()

            if not keys:
                raise ValueError("No datasets found in HDF5 file")

            target_key = key if key in keys else keys[0]

            if sample_rows:
                df = store.select(target_key, stop=sample_rows)
            else:
                df = store.select(target_key)

            metadata = {
                "keys": keys,
                "analyzed_key": target_key,
            }

        return df, metadata

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        key = kwargs.get("key")

        with pd.HDFStore(file_path, mode="r") as store:
            keys = store.keys()
            target_key = key if key in keys else keys[0]
            return store.get_storer(target_key).nrows


# ==================== Pickle Handler ====================

class PickleHandler(BaseFileHandler):
    """Handler for Python pickle files (DataFrame only)."""

    extensions = [".pkl", ".pickle"]
    file_type = "pickle"
    description = "Python Pickle (DataFrame)"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        obj = pd.read_pickle(file_path)

        if not isinstance(obj, pd.DataFrame):
            if isinstance(obj, pd.Series):
                obj = obj.to_frame()
            else:
                raise ValueError("Pickle file does not contain a DataFrame or Series")

        if sample_rows and len(obj) > sample_rows:
            obj = obj.head(sample_rows)

        return obj, {}

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        obj = pd.read_pickle(file_path)
        if isinstance(obj, pd.DataFrame):
            return len(obj)
        elif isinstance(obj, pd.Series):
            return len(obj)
        raise ValueError("Pickle file does not contain a DataFrame or Series")


# ==================== SAS Handler ====================

class SASHandler(BaseFileHandler):
    """Handler for SAS files (.sas7bdat, .xpt)."""

    extensions = [".sas7bdat", ".xpt"]
    file_type = "sas"
    description = "SAS Data File"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        ext = file_path.suffix.lower()

        if ext == ".xpt":
            df = pd.read_sas(file_path, format="xport")
        else:
            df = pd.read_sas(file_path)

        if sample_rows and len(df) > sample_rows:
            df = df.head(sample_rows)

        return df, {"sas_format": "xport" if ext == ".xpt" else "sas7bdat"}

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        df = pd.read_sas(file_path)
        return len(df)


# ==================== Stata Handler ====================

class StataHandler(BaseFileHandler):
    """Handler for Stata files (.dta)."""

    extensions = [".dta"]
    file_type = "stata"
    description = "Stata Data File"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        df = pd.read_stata(file_path)

        if sample_rows and len(df) > sample_rows:
            df = df.head(sample_rows)

        return df, {}

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        df = pd.read_stata(file_path)
        return len(df)


# ==================== SPSS Handler ====================

class SPSSHandler(BaseFileHandler):
    """Handler for SPSS files (.sav, .zsav)."""

    extensions = [".sav", ".zsav", ".por"]
    file_type = "spss"
    description = "SPSS Data File"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        df = pd.read_spss(file_path)

        if sample_rows and len(df) > sample_rows:
            df = df.head(sample_rows)

        return df, {}

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        df = pd.read_spss(file_path)
        return len(df)


# ==================== XML Handler ====================

class XMLHandler(BaseFileHandler):
    """Handler for XML files with tabular structure."""

    extensions = [".xml"]
    file_type = "xml"
    description = "XML (Extensible Markup Language)"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        xpath = kwargs.get("xpath", "//*[1]/*")

        df = pd.read_xml(file_path, xpath=xpath)

        if sample_rows and len(df) > sample_rows:
            df = df.head(sample_rows)

        return df, {"xpath": xpath}

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        xpath = kwargs.get("xpath", "//*[1]/*")
        df = pd.read_xml(file_path, xpath=xpath)
        return len(df)


# ==================== HTML Handler ====================

class HTMLHandler(BaseFileHandler):
    """Handler for HTML files with tables."""

    extensions = [".html", ".htm"]
    file_type = "html"
    description = "HTML Tables"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        table_index = kwargs.get("table_index", 0)

        tables = pd.read_html(file_path)

        if not tables:
            raise ValueError("No tables found in HTML file")

        if table_index >= len(tables):
            table_index = 0

        df = tables[table_index]

        if sample_rows and len(df) > sample_rows:
            df = df.head(sample_rows)

        metadata = {
            "table_count": len(tables),
            "analyzed_table_index": table_index,
        }
        return df, metadata

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        table_index = kwargs.get("table_index", 0)
        tables = pd.read_html(file_path)
        if table_index >= len(tables):
            table_index = 0
        return len(tables[table_index])


# ==================== ORC Handler ====================

class ORCHandler(BaseFileHandler):
    """Handler for Apache ORC files."""

    extensions = [".orc"]
    file_type = "orc"
    description = "Apache ORC (Optimized Row Columnar)"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        df = pd.read_orc(file_path)

        if sample_rows and len(df) > sample_rows:
            df = df.head(sample_rows)

        return df, {}

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        df = pd.read_orc(file_path)
        return len(df)


# ==================== Fixed Width Handler ====================

class FixedWidthHandler(BaseFileHandler):
    """Handler for fixed-width text files."""

    extensions = [".fwf", ".dat"]
    file_type = "fixed_width"
    description = "Fixed-Width Format"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        colspecs = kwargs.get("colspecs", "infer")

        df = pd.read_fwf(file_path, colspecs=colspecs, nrows=sample_rows)

        return df, {"colspecs": "inferred" if colspecs == "infer" else colspecs}

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        with open(file_path, "r") as f:
            return sum(1 for _ in f)


# ==================== Text Handler ====================

class TextHandler(BaseFileHandler):
    """Handler for plain text files (attempts to parse as delimited)."""

    extensions = [".txt", ".log"]
    file_type = "text"
    description = "Plain Text"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        lines = content.strip().split("\n")

        # Try to detect delimiter
        if lines:
            first_line = lines[0]
            for delimiter in [",", "\t", ";", "|", " "]:
                if delimiter in first_line:
                    try:
                        df = pd.read_csv(
                            file_path,
                            delimiter=delimiter,
                            nrows=sample_rows,
                            on_bad_lines="skip"
                        )
                        if len(df.columns) > 1:
                            return df, {"detected_delimiter": delimiter}
                    except Exception:
                        pass

        # Fall back to line-by-line
        sample_lines = lines[:sample_rows] if sample_rows else lines
        df = pd.DataFrame({
            "line_number": range(1, len(sample_lines) + 1),
            "content": sample_lines
        })

        return df, {"format": "plain_text"}

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        with open(file_path, "r", encoding="utf-8") as f:
            return sum(1 for _ in f)


# ==================== Word Document Handler ====================

class WordHandler(BaseFileHandler):
    """Handler for Word documents (.docx)."""

    extensions = [".docx", ".doc"]
    file_type = "word"
    description = "Microsoft Word Document"

    def read(
        self,
        file_path: Path,
        sample_rows: Optional[int] = None,
        **kwargs: Any
    ) -> tuple[pd.DataFrame, dict[str, Any]]:
        from docx import Document

        doc = Document(file_path)

        # Extract tables if any
        tables_data = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_rows.append(row_data)
            if table_rows:
                tables_data.append(table_rows)

        # If tables exist, use the first one
        if tables_data:
            first_table = tables_data[0]
            if len(first_table) > 1:
                headers = first_table[0]
                data = first_table[1:]
                df = pd.DataFrame(data, columns=headers)

                if sample_rows and len(df) > sample_rows:
                    df = df.head(sample_rows)

                return df, {
                    "has_tables": True,
                    "table_count": len(tables_data),
                }

        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        sample_paragraphs = paragraphs[:sample_rows] if sample_rows else paragraphs

        df = pd.DataFrame({
            "paragraph_number": range(1, len(sample_paragraphs) + 1),
            "content": sample_paragraphs
        })

        return df, {
            "has_tables": False,
            "paragraph_count": len(paragraphs),
        }

    def get_total_rows(self, file_path: Path, **kwargs: Any) -> int:
        from docx import Document

        doc = Document(file_path)

        # Check tables first
        for table in doc.tables:
            if len(table.rows) > 1:
                return len(table.rows) - 1  # Minus header

        # Fall back to paragraph count
        return len([p for p in doc.paragraphs if p.text.strip()])


# ==================== File Handler Registry ====================

class FileHandlerRegistry:
    """Registry for file type handlers."""

    def __init__(self):
        self._handlers: list[BaseFileHandler] = []
        self._extension_map: dict[str, BaseFileHandler] = {}
        self._register_default_handlers()

    def _register_default_handlers(self):
        """Register all default handlers."""
        default_handlers = [
            CSVHandler(),
            ExcelHandler(),
            JSONHandler(),
            ParquetHandler(),
            SQLiteHandler(),
            FeatherHandler(),
            HDF5Handler(),
            PickleHandler(),
            SASHandler(),
            StataHandler(),
            SPSSHandler(),
            XMLHandler(),
            HTMLHandler(),
            ORCHandler(),
            FixedWidthHandler(),
            TextHandler(),
            WordHandler(),
        ]

        for handler in default_handlers:
            self.register(handler)

    def register(self, handler: BaseFileHandler):
        """Register a file handler."""
        self._handlers.append(handler)
        for ext in handler.extensions:
            self._extension_map[ext.lower()] = handler

    def get_handler(
        self,
        file_path: Union[str, Path],
        content: Optional[bytes] = None
    ) -> Optional[BaseFileHandler]:
        """Get the appropriate handler for a file."""
        path = Path(file_path)
        ext = path.suffix.lower()

        # Try extension-based lookup first
        if ext in self._extension_map:
            return self._extension_map[ext]

        # Try content-based detection
        if content:
            for handler in self._handlers:
                if handler.can_handle(path, content):
                    return handler

        return None

    def get_supported_extensions(self) -> dict[str, str]:
        """Get all supported extensions with descriptions."""
        result = {}
        for handler in self._handlers:
            for ext in handler.extensions:
                result[ext] = handler.description
        return result

    def is_supported(self, file_path: Union[str, Path]) -> bool:
        """Check if a file extension is supported."""
        path = Path(file_path)
        ext = path.suffix.lower()
        return ext in self._extension_map


# Global registry instance
_registry = FileHandlerRegistry()


def get_handler(
    file_path: Union[str, Path],
    content: Optional[bytes] = None
) -> Optional[BaseFileHandler]:
    """Get handler for a file."""
    return _registry.get_handler(file_path, content)


def get_supported_extensions() -> dict[str, str]:
    """Get all supported file extensions."""
    return _registry.get_supported_extensions()


def is_supported(file_path: Union[str, Path]) -> bool:
    """Check if a file is supported."""
    return _registry.is_supported(file_path)


def read_file(
    file_path: Union[str, Path],
    sample_rows: Optional[int] = None,
    **kwargs: Any
) -> tuple[pd.DataFrame, dict[str, Any]]:
    """Read a file using the appropriate handler.

    Args:
        file_path: Path to the file
        sample_rows: Maximum rows to read (None for all)
        **kwargs: Handler-specific options

    Returns:
        Tuple of (DataFrame, metadata dict)

    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    handler = get_handler(path)
    if not handler:
        supported = ", ".join(get_supported_extensions().keys())
        raise ValueError(
            f"Unsupported file type: {path.suffix}. Supported: {supported}"
        )

    return handler.read(path, sample_rows=sample_rows, **kwargs)


def get_total_rows(file_path: Union[str, Path], **kwargs: Any) -> int:
    """Get total row count for a file.

    Args:
        file_path: Path to the file
        **kwargs: Handler-specific options

    Returns:
        Total number of rows
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    handler = get_handler(path)
    if not handler:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    return handler.get_total_rows(path, **kwargs)


def get_file_type(file_path: Union[str, Path]) -> str:
    """Get the file type name for a file.

    Args:
        file_path: Path to the file

    Returns:
        File type string (e.g., 'csv', 'excel', 'parquet')
    """
    path = Path(file_path)
    handler = get_handler(path)

    if not handler:
        supported = ", ".join(get_supported_extensions().keys())
        raise ValueError(
            f"Unsupported file type: {path.suffix}. Supported: {supported}"
        )

    return handler.file_type
