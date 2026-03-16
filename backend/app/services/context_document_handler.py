"""Context Document Handler - saves files and extracts text content.

This service handles:
- Saving uploaded context documents to disk
- Extracting text content from various file types
- Creating ContextDocument database records

Supported file types:
- PDF (.pdf) - using pypdf
- Word (.docx, .doc) - using python-docx
- Excel (.xlsx, .xls) - using pandas
- CSV (.csv) - using pandas
- Text (.txt, .md) - direct read
- HTML (.html, .htm) - using BeautifulSoup
- Images (.png, .jpg, .jpeg, .gif, .webp) - explanation only (no OCR)
"""
import logging
import os
import uuid
import re
from datetime import datetime
from pathlib import Path
from typing import BinaryIO, Optional

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.context_document import ContextDocument

logger = logging.getLogger(__name__)


# Supported file extensions by category
DOCUMENT_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md', '.html', '.htm'}
SPREADSHEET_EXTENSIONS = {'.xlsx', '.xls', '.csv'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.tiff'}
ALL_SUPPORTED_EXTENSIONS = DOCUMENT_EXTENSIONS | SPREADSHEET_EXTENSIONS | IMAGE_EXTENSIONS

# No file size limit for context documents (removed 10 MB restriction)


class ContextDocumentHandler:
    """Handles context document uploads and text extraction."""

    def __init__(self, db: Session):
        self.db = db
        self.settings = get_settings()
        self._ensure_upload_dir()

    def _ensure_upload_dir(self) -> Path:
        """Ensure the context documents upload directory exists."""
        context_dir = Path(self.settings.upload_dir) / "context_documents"
        context_dir.mkdir(parents=True, exist_ok=True)
        return context_dir

    def _get_file_type(self, extension: str) -> str:
        """Map file extension to file type."""
        ext = extension.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext in {'.docx', '.doc'}:
            return 'docx'
        elif ext in {'.xlsx', '.xls'}:
            return 'excel'
        elif ext == '.csv':
            return 'csv'
        elif ext == '.txt':
            return 'txt'
        elif ext == '.md':
            return 'md'
        elif ext in {'.html', '.htm'}:
            return 'html'
        elif ext in IMAGE_EXTENSIONS:
            return 'image'
        else:
            return 'unknown'

    def save_and_create(
        self,
        file: BinaryIO,
        original_filename: str,
        project_id: uuid.UUID,
        name: str,
        explanation: str,
    ) -> ContextDocument:
        """Save file to disk and create database record.

        Args:
            file: File-like object with the uploaded content
            original_filename: Original name of the uploaded file
            project_id: ID of the project this document belongs to
            name: User-provided name for the document
            explanation: User's explanation of what this document contains

        Returns:
            ContextDocument record

        Raises:
            ValueError: If file type is not supported or file is too large
        """
        # Validate extension
        extension = Path(original_filename).suffix.lower()
        if extension not in ALL_SUPPORTED_EXTENSIONS:
            supported = ', '.join(sorted(ALL_SUPPORTED_EXTENSIONS))
            raise ValueError(
                f"Unsupported file type: {extension}. Supported types: {supported}"
            )

        # Read file content
        content = file.read()
        file_size = len(content)

        # Generate unique filename
        doc_id = uuid.uuid4()
        safe_filename = f"{doc_id}{extension}"
        upload_dir = self._ensure_upload_dir()
        file_path = upload_dir / safe_filename

        # Save file to disk
        with open(file_path, 'wb') as f:
            f.write(content)

        # Get file type
        file_type = self._get_file_type(extension)

        # Create database record
        context_doc = ContextDocument(
            id=doc_id,
            project_id=project_id,
            name=name,
            original_filename=original_filename,
            file_path=str(file_path),
            file_type=file_type,
            file_size_bytes=file_size,
            explanation=explanation,
            extraction_status='pending',
        )

        self.db.add(context_doc)
        self.db.commit()
        self.db.refresh(context_doc)

        # Attempt text extraction (non-blocking, updates status)
        self._extract_text(context_doc)

        return context_doc

    def _extract_text(self, context_doc: ContextDocument) -> None:
        """Extract text content from the document.

        Updates the context_doc record with extracted text or error.
        """
        try:
            file_path = Path(context_doc.file_path)
            file_type = context_doc.file_type

            if file_type == 'image':
                # Images don't have extractable text - use explanation only
                context_doc.extracted_text = None
                context_doc.extraction_status = 'completed'
                logger.info(f"Image document {context_doc.id} - using explanation only")

            elif file_type == 'pdf':
                text = self._extract_pdf_text(file_path)
                context_doc.extracted_text = text
                context_doc.extraction_status = 'completed'

            elif file_type == 'docx':
                text = self._extract_docx_text(file_path)
                context_doc.extracted_text = text
                context_doc.extraction_status = 'completed'

            elif file_type in {'txt', 'md'}:
                text = self._extract_plain_text(file_path)
                context_doc.extracted_text = text
                context_doc.extraction_status = 'completed'

            elif file_type == 'html':
                text = self._extract_html_text(file_path)
                context_doc.extracted_text = text
                context_doc.extraction_status = 'completed'

            elif file_type == 'excel':
                text = self._extract_excel_text(file_path)
                context_doc.extracted_text = text
                context_doc.extraction_status = 'completed'

            elif file_type == 'csv':
                text = self._extract_csv_text(file_path)
                context_doc.extracted_text = text
                context_doc.extraction_status = 'completed'

            else:
                context_doc.extraction_status = 'failed'
                context_doc.extraction_error = f"Unsupported file type: {file_type}"

            self.db.commit()

        except Exception as e:
            logger.exception(f"Error extracting text from {context_doc.id}")
            context_doc.extraction_status = 'failed'
            context_doc.extraction_error = str(e)[:500]  # Truncate long errors
            self.db.commit()

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text_parts = []

        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")

        full_text = "\n\n".join(text_parts)

        # Clean up excessive whitespace
        full_text = re.sub(r'\n{3,}', '\n\n', full_text)
        full_text = re.sub(r' {2,}', ' ', full_text)

        return full_text.strip()

    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from Word document."""
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("\n[Table]\n" + "\n".join(table_text))

        return "\n\n".join(text_parts).strip()

    def _extract_plain_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except UnicodeDecodeError:
            # Try with latin-1 fallback
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read().strip()

    def _extract_html_text(self, file_path: Path) -> str:
        """Extract text from HTML file, stripping tags."""
        from bs4 import BeautifulSoup

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()

        soup = BeautifulSoup(content, 'html.parser')

        # Remove script and style elements
        for script in soup(['script', 'style', 'nav', 'footer', 'header']):
            script.decompose()

        # Get text
        text = soup.get_text(separator='\n')

        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines()]
        text = '\n'.join(line for line in lines if line)

        return text.strip()

    def _extract_excel_text(self, file_path: Path) -> str:
        """Extract text from Excel file (.xlsx, .xls).

        Converts each sheet to a readable text format with column headers and data.
        Useful for data dictionaries, feature lists, and tabular documentation.
        """
        import pandas as pd

        text_parts = []

        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names

        for sheet_name in sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            if df.empty:
                continue

            # Add sheet header
            text_parts.append(f"=== Sheet: {sheet_name} ===")

            # Add column info
            text_parts.append(f"Columns ({len(df.columns)}): {', '.join(str(c) for c in df.columns)}")
            text_parts.append(f"Rows: {len(df)}")
            text_parts.append("")

            # Convert to readable format (limit rows for very large files)
            max_rows = 500
            if len(df) > max_rows:
                # Show first and last rows
                preview_df = pd.concat([df.head(250), df.tail(250)])
                text_parts.append(f"[Showing first 250 and last 250 of {len(df)} rows]")
                text_parts.append(preview_df.to_string(index=False))
            else:
                text_parts.append(df.to_string(index=False))

            text_parts.append("")

        return "\n".join(text_parts).strip()

    def _extract_csv_text(self, file_path: Path) -> str:
        """Extract text from CSV file.

        Converts CSV to a readable text format with column headers and data.
        Useful for data dictionaries, feature lists, and tabular documentation.
        """
        import pandas as pd

        text_parts = []

        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode CSV file with any supported encoding")

        if df.empty:
            return "[Empty CSV file]"

        # Add column info
        text_parts.append(f"Columns ({len(df.columns)}): {', '.join(str(c) for c in df.columns)}")
        text_parts.append(f"Rows: {len(df)}")
        text_parts.append("")

        # Convert to readable format (limit rows for very large files)
        max_rows = 500
        if len(df) > max_rows:
            # Show first and last rows
            preview_df = pd.concat([df.head(250), df.tail(250)])
            text_parts.append(f"[Showing first 250 and last 250 of {len(df)} rows]")
            text_parts.append(preview_df.to_string(index=False))
        else:
            text_parts.append(df.to_string(index=False))

        return "\n".join(text_parts).strip()

    def re_extract_text(self, context_doc: ContextDocument) -> None:
        """Re-run text extraction for a document.

        Useful if extraction initially failed or to refresh content.
        """
        context_doc.extraction_status = 'pending'
        context_doc.extraction_error = None
        context_doc.extracted_text = None
        self.db.commit()

        self._extract_text(context_doc)

    def delete_document(self, context_doc: ContextDocument) -> None:
        """Delete a context document and its file."""
        # Delete file from disk
        try:
            file_path = Path(context_doc.file_path)
            if file_path.exists():
                file_path.unlink()
        except Exception as e:
            logger.warning(f"Failed to delete file {context_doc.file_path}: {e}")

        # Delete database record
        self.db.delete(context_doc)
        self.db.commit()


def get_supported_extensions() -> dict[str, str]:
    """Get dict of supported extensions with descriptions."""
    return {
        '.pdf': 'PDF Document',
        '.docx': 'Microsoft Word Document',
        '.doc': 'Microsoft Word Document (Legacy)',
        '.xlsx': 'Microsoft Excel Spreadsheet',
        '.xls': 'Microsoft Excel Spreadsheet (Legacy)',
        '.csv': 'CSV File',
        '.txt': 'Plain Text',
        '.md': 'Markdown',
        '.html': 'HTML Document',
        '.htm': 'HTML Document',
        '.png': 'PNG Image',
        '.jpg': 'JPEG Image',
        '.jpeg': 'JPEG Image',
        '.gif': 'GIF Image',
        '.webp': 'WebP Image',
        '.bmp': 'Bitmap Image',
        '.tiff': 'TIFF Image',
    }
