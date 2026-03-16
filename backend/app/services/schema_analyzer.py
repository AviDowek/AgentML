"""Schema analysis service for data sources."""
import logging
import pandas as pd
from pathlib import Path
from typing import Any

from app.services.file_handlers import (
    get_supported_extensions,
    get_file_type as _handler_get_file_type,
    read_file,
    get_total_rows,
    get_handler,
    is_supported,
)

logger = logging.getLogger(__name__)

# Build SUPPORTED_EXTENSIONS from the file handlers registry for backward compatibility
SUPPORTED_EXTENSIONS = get_supported_extensions()


def get_file_type(file_path: str | Path) -> str:
    """Get the file type from extension.

    Args:
        file_path: Path to the file

    Returns:
        File type string

    Raises:
        ValueError: If file type is not supported
    """
    return _handler_get_file_type(file_path)


class SchemaAnalyzer:
    """Analyzes data files to extract schema and statistics."""

    def __init__(self, sample_rows: int = 10000):
        """Initialize analyzer.

        Args:
            sample_rows: Number of rows to sample for analysis
        """
        self.sample_rows = sample_rows

    def analyze_file(self, file_path: str | Path, **kwargs) -> dict[str, Any]:
        """Analyze any supported file and return schema summary.

        Uses the extensible file handler system to support many file formats.

        Args:
            file_path: Path to the file
            **kwargs: Additional arguments (delimiter for CSV, sheet_name for Excel, etc.)

        Returns:
            Dictionary containing schema information and statistics
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Get handler for this file type
        handler = get_handler(file_path)
        if not handler:
            supported = ", ".join(SUPPORTED_EXTENSIONS.keys())
            raise ValueError(f"Unsupported file type: {file_path.suffix}. Supported: {supported}")

        file_type = handler.file_type

        # For legacy file types, use existing methods for backward compatibility
        if file_type == "csv":
            return self.analyze_csv(file_path, delimiter=kwargs.get("delimiter", ","))
        elif file_type == "excel":
            return self.analyze_excel(file_path, sheet_name=kwargs.get("sheet_name", 0))
        elif file_type == "json":
            return self.analyze_json(file_path)
        elif file_type == "parquet":
            return self.analyze_parquet(file_path)
        elif file_type == "text":
            return self.analyze_text(file_path)
        elif file_type == "word":
            return self.analyze_word(file_path)
        else:
            # Use the new handler system for all other file types
            return self._analyze_with_handler(file_path, handler, **kwargs)

    def _analyze_with_handler(self, file_path: Path, handler, **kwargs) -> dict[str, Any]:
        """Analyze a file using the handler-based system.

        Args:
            file_path: Path to the file
            handler: File handler instance
            **kwargs: Handler-specific options

        Returns:
            Dictionary containing schema information and statistics
        """
        try:
            # Read file with sampling
            df, metadata = handler.read(file_path, sample_rows=self.sample_rows, **kwargs)

            # Get total row count
            try:
                total_rows = handler.get_total_rows(file_path, **kwargs)
            except Exception:
                total_rows = len(df)

            # Analyze the DataFrame
            result = self._analyze_dataframe(df, total_rows)
            result["file_type"] = handler.file_type

            # Add any handler-specific metadata
            if metadata:
                result.update(metadata)

            return result

        except Exception as e:
            logger.error(f"Error analyzing file {file_path} with handler {handler.file_type}: {e}")
            raise

    def analyze_csv(self, file_path: str | Path, delimiter: str = ",") -> dict[str, Any]:
        """Analyze a CSV file and return schema summary.

        Args:
            file_path: Path to the CSV file
            delimiter: CSV delimiter character

        Returns:
            Dictionary containing schema information and statistics
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Read sample of the file
        df = pd.read_csv(file_path, delimiter=delimiter, nrows=self.sample_rows)

        # Get total row count (read just the index)
        with open(file_path, "r", encoding="utf-8") as f:
            total_rows = sum(1 for _ in f) - 1  # -1 for header

        result = self._analyze_dataframe(df, total_rows)
        result["file_type"] = "csv"
        return result

    def analyze_excel(self, file_path: str | Path, sheet_name: str | int = 0) -> dict[str, Any]:
        """Analyze an Excel file and return schema summary.

        Args:
            file_path: Path to the Excel file
            sheet_name: Sheet name or index to analyze

        Returns:
            Dictionary containing schema information and statistics
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Read the Excel file with context manager to ensure it's closed
        with pd.ExcelFile(file_path) as xlsx:
            sheet_names = xlsx.sheet_names

            # Read sample of the specified sheet
            df = pd.read_excel(xlsx, sheet_name=sheet_name, nrows=self.sample_rows)

            # Get total row count by reading without limit
            df_full = pd.read_excel(xlsx, sheet_name=sheet_name)
            total_rows = len(df_full)

        result = self._analyze_dataframe(df, total_rows)
        result["file_type"] = "excel"
        result["sheet_names"] = sheet_names
        result["analyzed_sheet"] = sheet_name if isinstance(sheet_name, str) else sheet_names[sheet_name]
        return result

    def analyze_json(self, file_path: str | Path) -> dict[str, Any]:
        """Analyze a JSON file and return schema summary.

        Args:
            file_path: Path to the JSON file

        Returns:
            Dictionary containing schema information and statistics
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Try to read as JSON array or object
        df = pd.read_json(file_path)

        # Sample if too large
        if len(df) > self.sample_rows:
            df_sample = df.head(self.sample_rows)
            total_rows = len(df)
        else:
            df_sample = df
            total_rows = len(df)

        result = self._analyze_dataframe(df_sample, total_rows)
        result["file_type"] = "json"
        return result

    def analyze_parquet(self, file_path: str | Path) -> dict[str, Any]:
        """Analyze a Parquet file and return schema summary.

        Args:
            file_path: Path to the Parquet file

        Returns:
            Dictionary containing schema information and statistics
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Read parquet file
        df = pd.read_parquet(file_path)

        # Sample if too large
        if len(df) > self.sample_rows:
            df_sample = df.head(self.sample_rows)
            total_rows = len(df)
        else:
            df_sample = df
            total_rows = len(df)

        result = self._analyze_dataframe(df_sample, total_rows)
        result["file_type"] = "parquet"
        return result

    def analyze_text(self, file_path: str | Path) -> dict[str, Any]:
        """Analyze a text file and return content summary.

        For .txt files, we extract lines as rows with a single 'content' column.
        If it looks like a delimited file, we try to parse it.

        Args:
            file_path: Path to the text file

        Returns:
            Dictionary containing content information
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        lines = content.strip().split("\n")
        total_lines = len(lines)

        # Check if it might be a delimited file (CSV-like)
        if total_lines > 1:
            first_line = lines[0]
            # Check for common delimiters
            for delimiter in [",", "\t", ";", "|"]:
                if delimiter in first_line:
                    # Try to parse as delimited
                    try:
                        df = pd.read_csv(file_path, delimiter=delimiter, nrows=self.sample_rows)
                        if len(df.columns) > 1:
                            with open(file_path, "r", encoding="utf-8") as f:
                                row_count = sum(1 for _ in f) - 1
                            result = self._analyze_dataframe(df, row_count)
                            result["file_type"] = "text"
                            result["detected_format"] = f"delimited ({delimiter})"
                            return result
                    except Exception:
                        pass

        # Treat as plain text - create a DataFrame with line content
        sample_lines = lines[:self.sample_rows]
        df = pd.DataFrame({"line_number": range(1, len(sample_lines) + 1), "content": sample_lines})

        # Calculate text statistics
        char_count = len(content)
        word_count = len(content.split())
        avg_line_length = char_count / total_lines if total_lines > 0 else 0

        return {
            "file_type": "text",
            "row_count": total_lines,
            "column_count": 2,
            "sample_rows": len(sample_lines),
            "columns": [
                {
                    "name": "line_number",
                    "dtype": "int64",
                    "inferred_type": "numeric",
                    "null_count": 0,
                    "null_percentage": 0.0,
                    "unique_count": len(sample_lines),
                },
                {
                    "name": "content",
                    "dtype": "object",
                    "inferred_type": "text",
                    "null_count": 0,
                    "null_percentage": 0.0,
                    "unique_count": len(set(sample_lines)),
                },
            ],
            "text_stats": {
                "total_lines": total_lines,
                "total_characters": char_count,
                "total_words": word_count,
                "avg_line_length": round(avg_line_length, 2),
            },
        }

    def analyze_word(self, file_path: str | Path) -> dict[str, Any]:
        """Analyze a Word document and return content summary.

        Args:
            file_path: Path to the Word document

        Returns:
            Dictionary containing content information
        """
        from docx import Document

        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        doc = Document(file_path)

        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract tables if any
        tables_data = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_rows.append(row_data)
            if table_rows:
                tables_data.append(table_rows)

        # If there are tables, analyze the first one as a DataFrame
        if tables_data:
            first_table = tables_data[0]
            if len(first_table) > 1:
                # Use first row as header
                headers = first_table[0]
                data = first_table[1:]
                df = pd.DataFrame(data, columns=headers)

                result = self._analyze_dataframe(df, len(data))
                result["file_type"] = "word"
                result["has_tables"] = True
                result["table_count"] = len(tables_data)
                result["paragraph_count"] = len(paragraphs)
                return result

        # No tables - treat as text content
        total_paragraphs = len(paragraphs)
        full_text = "\n".join(paragraphs)
        char_count = len(full_text)
        word_count = len(full_text.split())

        sample_paragraphs = paragraphs[:self.sample_rows]

        return {
            "file_type": "word",
            "row_count": total_paragraphs,
            "column_count": 2,
            "sample_rows": len(sample_paragraphs),
            "columns": [
                {
                    "name": "paragraph_number",
                    "dtype": "int64",
                    "inferred_type": "numeric",
                    "null_count": 0,
                    "null_percentage": 0.0,
                    "unique_count": len(sample_paragraphs),
                },
                {
                    "name": "content",
                    "dtype": "object",
                    "inferred_type": "text",
                    "null_count": 0,
                    "null_percentage": 0.0,
                    "unique_count": len(set(sample_paragraphs)),
                },
            ],
            "text_stats": {
                "total_paragraphs": total_paragraphs,
                "total_characters": char_count,
                "total_words": word_count,
                "has_tables": len(tables_data) > 0,
                "table_count": len(tables_data),
            },
        }

    def analyze_dataframe(self, df: pd.DataFrame) -> dict[str, Any]:
        """Analyze a DataFrame and return schema summary.

        Args:
            df: Pandas DataFrame to analyze

        Returns:
            Dictionary containing schema information and statistics
        """
        return self._analyze_dataframe(df, len(df))

    def _analyze_dataframe(self, df: pd.DataFrame, total_rows: int) -> dict[str, Any]:
        """Internal method to analyze DataFrame.

        Args:
            df: Pandas DataFrame (possibly sampled)
            total_rows: Total number of rows in full dataset

        Returns:
            Dictionary containing schema information and statistics
        """
        columns = []

        for col in df.columns:
            col_info = self._analyze_column(df[col])
            col_info["name"] = str(col)  # Ensure column name is string
            columns.append(col_info)

        result = {
            "row_count": total_rows,
            "column_count": len(df.columns),
            "sample_rows": len(df),
            "columns": columns,
        }

        # Check if this looks like a data dictionary/metadata file
        warnings = self._detect_data_dictionary(df, columns)
        if warnings:
            result["warnings"] = warnings

        return result

    def _detect_data_dictionary(self, df: pd.DataFrame, columns: list[dict[str, Any]]) -> list[str]:
        """Detect if a file looks like a data dictionary rather than actual data.

        Args:
            df: The DataFrame to analyze
            columns: The analyzed column info

        Returns:
            List of warning messages if issues detected, empty list otherwise
        """
        warnings = []

        # Check for auto-generated "Unnamed" columns (common when no proper header)
        unnamed_cols = [c for c in df.columns if str(c).startswith("Unnamed:")]
        if len(unnamed_cols) > 0:
            unnamed_ratio = len(unnamed_cols) / len(df.columns)
            if unnamed_ratio > 0.3:  # More than 30% unnamed columns
                warnings.append(
                    f"This file has {len(unnamed_cols)} auto-generated column names "
                    f"('Unnamed: 0', etc.), suggesting it may not have proper headers. "
                    "This could indicate a metadata document or improperly formatted data."
                )

        # Check for metadata-like content patterns
        metadata_keywords = [
            "description", "field name", "field type", "data type", "column name",
            "definition", "format", "notes", "limitations", "idiosyncrasies",
            "agency", "category", "documentation", "metadata"
        ]

        # Check column names for metadata keywords
        col_names_lower = [str(c).lower() for c in df.columns]
        metadata_col_matches = sum(
            1 for name in col_names_lower
            for keyword in metadata_keywords
            if keyword in name
        )

        # Check cell content for metadata keywords (sample first few cells)
        if len(df) > 0:
            text_cols = [c for c in df.columns if df[c].dtype == "object"]
            metadata_content_count = 0
            for col in text_cols[:5]:  # Check first 5 text columns
                sample_values = df[col].dropna().head(20).astype(str).str.lower()
                for val in sample_values:
                    if any(keyword in val for keyword in metadata_keywords):
                        metadata_content_count += 1

            if metadata_content_count > 5 or metadata_col_matches > 2:
                warnings.append(
                    "This file appears to contain metadata field descriptions "
                    "(like 'description', 'field name', 'data type') rather than actual data. "
                    "If you intended to upload a data dictionary, visualizations may not be meaningful. "
                    "Please ensure you're uploading the actual dataset file."
                )

        # Check for very low row count with long text content (common in data dictionaries)
        if total_rows := len(df):
            if total_rows < 50:
                text_cols = [c for c in df.columns if df[c].dtype == "object"]
                if text_cols:
                    avg_text_length = df[text_cols].astype(str).apply(lambda x: x.str.len().mean()).mean()
                    if avg_text_length > 100:  # Average text > 100 chars with few rows
                        if "metadata" not in " ".join(warnings).lower():  # Avoid duplicate warning
                            warnings.append(
                                f"This file has only {total_rows} rows with long text content "
                                "(avg length > 100 chars), which is typical of data dictionaries "
                                "or documentation files rather than tabular datasets."
                            )

        return warnings

    def _analyze_column(self, series: pd.Series) -> dict[str, Any]:
        """Analyze a single column.

        Args:
            series: Pandas Series to analyze

        Returns:
            Dictionary containing column statistics
        """
        # Determine inferred type
        dtype = str(series.dtype)
        inferred_type = self._infer_semantic_type(series, dtype)

        # Calculate unique count safely (handles unhashable types like dicts/lists)
        try:
            unique_count = int(series.nunique())
        except TypeError:
            # Column contains unhashable types (dict, list, etc.)
            # Convert to strings first to count unique values
            try:
                unique_count = int(series.astype(str).nunique())
            except Exception:
                unique_count = 0

        stats: dict[str, Any] = {
            "dtype": dtype,
            "inferred_type": inferred_type,
            "null_count": int(series.isnull().sum()),
            "null_percentage": round(series.isnull().mean() * 100, 2),
            "unique_count": unique_count,
        }

        # Add type-specific statistics
        if inferred_type == "numeric":
            stats.update(self._numeric_stats(series))
        elif inferred_type == "categorical":
            stats.update(self._categorical_stats(series))
        elif inferred_type == "datetime":
            stats.update(self._datetime_stats(series))

        return stats

    def _infer_semantic_type(self, series: pd.Series, dtype: str) -> str:
        """Infer the semantic type of a column.

        Args:
            series: Pandas Series
            dtype: Pandas dtype string

        Returns:
            Semantic type: 'numeric', 'categorical', 'datetime', 'text', or 'boolean'
        """
        # Check for boolean
        if dtype == "bool":
            return "boolean"
        try:
            unique_vals = set(series.dropna().unique())
            if unique_vals.issubset({True, False, 0, 1}):
                return "boolean"
        except TypeError:
            # Column contains unhashable types (dict, list, etc.)
            pass

        # Check for numeric types
        if dtype in ("int64", "int32", "float64", "float32", "Int64", "Float64"):
            return "numeric"

        # Check for datetime
        if "datetime" in dtype:
            return "datetime"

        # Try to parse as datetime
        if dtype == "object":
            try:
                pd.to_datetime(series.dropna().head(100))
                return "datetime"
            except (ValueError, TypeError):
                pass

        # Check if categorical (low cardinality string)
        if dtype == "object":
            try:
                nunique = series.nunique()
                unique_ratio = nunique / len(series) if len(series) > 0 else 0
                if unique_ratio < 0.05 or nunique < 50:
                    return "categorical"
                return "text"
            except TypeError:
                # Column contains unhashable types (dict, list, etc.)
                # Treat as text/complex type
                return "text"

        return "unknown"

    def _numeric_stats(self, series: pd.Series) -> dict[str, Any]:
        """Calculate numeric column statistics."""
        clean = series.dropna()
        if len(clean) == 0:
            return {}

        return {
            "min": float(clean.min()),
            "max": float(clean.max()),
            "mean": round(float(clean.mean()), 4),
            "median": float(clean.median()),
            "std": round(float(clean.std()), 4) if len(clean) > 1 else 0,
        }

    def _categorical_stats(self, series: pd.Series) -> dict[str, Any]:
        """Calculate categorical column statistics."""
        try:
            value_counts = series.value_counts()
            top_values = value_counts.head(10).to_dict()

            return {
                "top_values": {str(k): int(v) for k, v in top_values.items()},
                "mode": str(value_counts.index[0]) if len(value_counts) > 0 else None,
            }
        except TypeError:
            # Column contains unhashable types (dict, list, etc.)
            # Convert to strings first
            try:
                str_series = series.astype(str)
                value_counts = str_series.value_counts()
                top_values = value_counts.head(10).to_dict()
                return {
                    "top_values": {str(k): int(v) for k, v in top_values.items()},
                    "mode": str(value_counts.index[0]) if len(value_counts) > 0 else None,
                }
            except Exception:
                return {"top_values": {}, "mode": None}

    def _datetime_stats(self, series: pd.Series) -> dict[str, Any]:
        """Calculate datetime column statistics."""
        try:
            dt_series = pd.to_datetime(series.dropna())
            if len(dt_series) == 0:
                return {}

            return {
                "min": str(dt_series.min()),
                "max": str(dt_series.max()),
            }
        except (ValueError, TypeError):
            return {}
